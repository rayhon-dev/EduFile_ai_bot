import os
import fitz  # PyMuPDF
from PIL import Image
import pytesseract
from docx import Document
from lxml import etree
import zipfile

from utils.translator import translate_text_preserving_math

def extract_equations_from_docx(docx_path):
    """DOCX ichidan OMML formulalarni chiqaradi."""
    equations = []
    try:
        with zipfile.ZipFile(docx_path) as docx_zip:
            for name in docx_zip.namelist():
                if name.startswith("word/") and name.endswith(".xml"):
                    xml_content = docx_zip.read(name)
                    try:
                        tree = etree.fromstring(xml_content)
                        ns = {"m": "http://schemas.openxmlformats.org/officeDocument/2006/math"}
                        for eq in tree.findall(".//m:oMath", ns):
                            eq_text = "".join(eq.itertext())
                            if eq_text.strip():
                                equations.append(eq_text.strip())
                    except Exception:
                        continue
    except Exception as e:
        print(f"[Equation extraction failed] {e}")
    return equations

def read_pdf_smart(file_path):
    """PyMuPDF orqali PDF o‘qiydi, agar text bo‘lmasa OCR bajaradi."""
    text = ""
    try:
        doc = fitz.open(file_path)
        for page in doc:
            page_text = page.get_text("text")
            if page_text.strip():
                text += page_text + "\n"
            else:
                pix = page.get_pixmap()
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                text += pytesseract.image_to_string(img, lang="eng+uz+rus") + "\n"
    except Exception as e:
        print(f"[PDF read error] {e}")
    return text.strip()

def read_file_content(file_path: str) -> str:
    """.txt, .md, .pdf, .docx fayllardan matn o‘qiydi."""
    try:
        ext = os.path.splitext(file_path)[1].lower()

        if ext in (".txt", ".md"):
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()

        elif ext == ".pdf":
            return read_pdf_smart(file_path)

        elif ext == ".docx":
            doc = Document(file_path)
            text_parts = []

            for p in doc.paragraphs:
                if p.text.strip():
                    text_parts.append(p.text.strip())

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text and cell_text not in text_parts:
                            text_parts.append(cell_text)

            equations = extract_equations_from_docx(file_path)
            if equations:
                text_parts.append("\n🧮 Extracted Formulas:\n" + "\n".join(equations))

            return "\n".join(text_parts).strip()

        else:
            return "❗ Fayl turi qo‘llab-quvvatlanmaydi. Faqat .txt, .md, .pdf, .docx ishlaydi."

    except Exception as e:
        print(f"[File read error] {e}")
        return None

def process_file_for_translation(file_path, source_lang="uz", target_lang="en"):
    """
    Faylni o‘qib, matematik ifodalarni saqlab tarjima qiladi.
    """
    text = read_file_content(file_path)
    if not text:
        return None

    # Tarjima uchun quyidagi funksiyaga faqat text beriladi,
    # matematik ifodalar maskalanadi va tarjimadan keyin tiklanadi.
    translated_text = translate_text_preserving_math(text, source_lang=source_lang, target_lang=target_lang)
    return translated_text
